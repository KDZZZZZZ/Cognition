import asyncio
import base64
import uuid
import re
from typing import List
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument

from app.config import settings
from app.models import DocumentChunk
from app.services.llm_service import llm_service


class DocumentParser:
    """Parse various document formats and extract structured content."""

    PDF_TARGET_CHUNK_CHARS = 1200
    PDF_MAX_CHUNK_CHARS = 2400

    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def parse_file(
        self,
        file_path: str,
        file_id: str,
        file_type: str
    ) -> tuple[List[DocumentChunk], dict]:
        """
        Parse a file and return chunks with metadata.

        Returns:
            (chunks, metadata) - List of DocumentChunk objects and file metadata
        """
        path = Path(file_path)

        if file_type == "pdf":
            return await self._parse_pdf(path, file_id)
        elif file_type in ["docx", "doc"]:
            return await self._parse_docx(path, file_id)
        elif file_type == "md":
            return await self._parse_markdown(path, file_id)
        elif file_type == "txt":
            return await self._parse_text(path, file_id)
        elif file_type == "web":
            return await self._parse_web_html(path, file_id)
        else:
            return [], {}

    async def _parse_pdf(self, path: Path, file_id: str) -> tuple[List[DocumentChunk], dict]:
        """Parse PDF file with page layout information."""
        chunks = []
        metadata = {"page_count": 0, "ocr_pages": []}

        # Use pdfplumber for accurate layout and coordinates
        with pdfplumber.open(path) as pdf:
            metadata["page_count"] = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    words = page.extract_words(extra_attrs=["fontname", "size"])
                    paragraphs = self._group_words_to_paragraphs(words)
                    paragraphs = self._merge_pdf_paragraphs(paragraphs)

                    for idx, para in enumerate(paragraphs):
                        chunk = DocumentChunk(
                            id=str(uuid.uuid4()),
                            file_id=file_id,
                            page=page_num + 1,
                            chunk_index=idx,
                            content=para["text"],
                            bbox=para.get("bbox")
                        )
                        chunks.append(chunk)
                    continue

                ocr_text = await self._ocr_pdf_page(path, page_num + 1)
                if not ocr_text:
                    continue

                metadata["ocr_pages"].append(page_num + 1)
                ocr_chunks = self._chunk_plain_text(ocr_text)
                for idx, text in enumerate(ocr_chunks):
                    chunks.append(
                        DocumentChunk(
                            id=str(uuid.uuid4()),
                            file_id=file_id,
                            page=page_num + 1,
                            chunk_index=idx,
                            content=text,
                            bbox=None,
                        )
                    )

        return chunks, metadata

    async def _parse_docx(self, path: Path, file_id: str) -> tuple[List[DocumentChunk], dict]:
        """Parse DOCX file."""
        chunks = []
        doc = DocxDocument(path)

        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    file_id=file_id,
                    page=1,
                    chunk_index=idx,
                    content=para.text,
                    bbox=None
                )
                chunks.append(chunk)

        metadata = {
            "page_count": 1,
            "paragraph_count": len(doc.paragraphs)
        }

        return chunks, metadata

    async def _parse_markdown(self, path: Path, file_id: str) -> tuple[List[DocumentChunk], dict]:
        """Parse Markdown file."""
        content = path.read_text(encoding="utf-8")
        chunks = []

        # Split by headers to create chunks
        lines = content.split("\n")
        current_chunk = []
        chunk_idx = 0

        for line in lines:
            if line.startswith("#"):
                # Save previous chunk
                if current_chunk:
                    chunk = DocumentChunk(
                        id=str(uuid.uuid4()),
                        file_id=file_id,
                        page=1,
                        chunk_index=chunk_idx,
                        content="\n".join(current_chunk),
                        bbox=None
                    )
                    chunks.append(chunk)
                    chunk_idx += 1
                    current_chunk = [line]
                else:
                    current_chunk.append(line)
            else:
                current_chunk.append(line)

        # Don't forget the last chunk
        if current_chunk:
            chunk = DocumentChunk(
                id=str(uuid.uuid4()),
                file_id=file_id,
                page=1,
                chunk_index=chunk_idx,
                content="\n".join(current_chunk),
                bbox=None
            )
            chunks.append(chunk)

        metadata = {"page_count": 1, "chunk_count": len(chunks)}
        return chunks, metadata

    async def _parse_text(self, path: Path, file_id: str) -> tuple[List[DocumentChunk], dict]:
        """Parse plain text file."""
        content = path.read_text(encoding="utf-8")

        # Split into paragraphs
        paragraphs = content.split("\n\n")
        chunks = []

        for idx, para in enumerate(paragraphs):
            if para.strip():
                chunk = DocumentChunk(
                    id=str(uuid.uuid4()),
                    file_id=file_id,
                    page=1,
                    chunk_index=idx,
                    content=para.strip(),
                    bbox=None
                )
                chunks.append(chunk)

        metadata = {"page_count": 1, "chunk_count": len(chunks)}
        return chunks, metadata

    async def _parse_web_html(self, path: Path, file_id: str) -> tuple[List[DocumentChunk], dict]:
        """Parse HTML/web snapshot file into text chunks."""
        content = path.read_text(encoding="utf-8", errors="ignore")
        chunks: List[DocumentChunk] = []

        try:
            from bs4 import BeautifulSoup
        except Exception:
            # Fallback strip tags using regex when bs4 is unavailable.
            text = " ".join(content.replace("\n", " ").split())
            text = re.sub(r"<[^>]+>", " ", text)
            blocks = [blk.strip() for blk in text.split(". ") if blk.strip()]
        else:
            soup = BeautifulSoup(content, "html.parser")
            nodes = soup.select("h1, h2, h3, h4, p, li, pre, code, blockquote, table")
            blocks = []
            for node in nodes:
                text = " ".join(node.get_text(" ", strip=True).split())
                if text:
                    blocks.append(text)

        for idx, block in enumerate(blocks):
            chunk = DocumentChunk(
                id=str(uuid.uuid4()),
                file_id=file_id,
                page=1,
                chunk_index=idx,
                content=block,
                bbox=None,
            )
            chunks.append(chunk)

        metadata = {"page_count": 1, "chunk_count": len(chunks), "source_type": "web"}
        return chunks, metadata

    def _group_words_to_paragraphs(self, words: List[dict]) -> List[dict]:
        """Group words into paragraphs based on vertical proximity."""
        if not words:
            return []

        paragraphs = []
        current_para = []
        last_y = None
        y_tolerance = 5  # pixels

        for word in words:
            x0, top, x1, bottom = word.get("x0"), word.get("top"), word.get("x1"), word.get("bottom")

            if last_y is not None and abs(top - last_y) > y_tolerance:
                # New paragraph
                if current_para:
                    # Sort words by x position and join
                    sorted_words = sorted(current_para, key=lambda w: w["x0"])
                    text = " ".join(w["text"] for w in sorted_words)

                    # Calculate bounding box
                    all_x0 = [w["x0"] for w in current_para]
                    all_y0 = [w["top"] for w in current_para]
                    all_x1 = [w["x1"] for w in current_para]
                    all_y1 = [w["bottom"] for w in current_para]

                    paragraphs.append({
                        "text": text,
                        "bbox": (min(all_x0), min(all_y0), max(all_x1), max(all_y1))
                    })
                    current_para = []

            current_para.append(word)
            last_y = top

        # Handle remaining words
        if current_para:
            sorted_words = sorted(current_para, key=lambda w: w["x0"])
            text = " ".join(w["text"] for w in sorted_words)

            all_x0 = [w["x0"] for w in current_para]
            all_y0 = [w["top"] for w in current_para]
            all_x1 = [w["x1"] for w in current_para]
            all_y1 = [w["bottom"] for w in current_para]

            paragraphs.append({
                "text": text,
                "bbox": (min(all_x0), min(all_y0), max(all_x1), max(all_y1))
            })

        return paragraphs

    async def _ocr_pdf_page(self, path: Path, page_num: int) -> str:
        if not settings.OCR_ENABLED or not llm_service.supports_ocr():
            return ""

        try:
            data_url = await asyncio.to_thread(self._render_pdf_page_as_data_url, path, page_num)
        except Exception:
            return ""
        if not data_url:
            return ""

        try:
            text = await llm_service.ocr_image(
                data_url,
                prompt=(
                    "Extract all visible text from this PDF page image. "
                    "Return plain text only. Preserve formulas, tables, and section numbers when readable."
                ),
            )
        except Exception:
            return ""
        return " ".join(text.split())

    def _render_pdf_page_as_data_url(self, path: Path, page_num: int) -> str:
        import pypdfium2 as pdfium  # type: ignore

        document = pdfium.PdfDocument(str(path))
        try:
            if page_num < 1 or page_num > len(document):
                return ""
            page = document[page_num - 1]
            bitmap = page.render(scale=max(float(settings.OCR_RENDER_DPI) / 72.0, 1.0))
            pil_image = bitmap.to_pil()
            try:
                if pil_image.mode != "RGB":
                    pil_image = pil_image.convert("RGB")
                from io import BytesIO

                buffer = BytesIO()
                pil_image.save(buffer, format="JPEG", quality=82, optimize=True)
                encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
                return f"data:image/jpeg;base64,{encoded}"
            finally:
                page.close()
                bitmap.close()
        finally:
            document.close()

    def _chunk_plain_text(self, text: str) -> List[str]:
        blocks = [block.strip() for block in re.split(r"\n{2,}", text) if block.strip()]
        if not blocks:
            compact = " ".join(str(text or "").split())
            blocks = [compact] if compact else []

        chunks: List[str] = []
        current: List[str] = []
        current_chars = 0

        for block in blocks:
            next_chars = current_chars + len(block) + (2 if current else 0)
            if current and current_chars >= self.PDF_TARGET_CHUNK_CHARS and next_chars > self.PDF_MAX_CHUNK_CHARS:
                chunks.append("\n\n".join(current))
                current = []
                current_chars = 0

            current.append(block)
            current_chars += len(block) + (2 if current_chars else 0)
            if current_chars >= self.PDF_MAX_CHUNK_CHARS:
                chunks.append("\n\n".join(current))
                current = []
                current_chars = 0

        if current:
            chunks.append("\n\n".join(current))

        return chunks

    def _merge_pdf_paragraphs(self, paragraphs: List[dict]) -> List[dict]:
        """Merge short adjacent PDF paragraphs into larger semantic chunks."""
        if not paragraphs:
            return []

        merged: List[dict] = []
        current_texts: List[str] = []
        current_boxes: List[tuple[float, float, float, float]] = []
        current_chars = 0

        def flush() -> None:
            nonlocal current_texts, current_boxes, current_chars
            if not current_texts:
                return
            xs0 = [box[0] for box in current_boxes]
            ys0 = [box[1] for box in current_boxes]
            xs1 = [box[2] for box in current_boxes]
            ys1 = [box[3] for box in current_boxes]
            merged.append(
                {
                    "text": "\n".join(current_texts),
                    "bbox": (min(xs0), min(ys0), max(xs1), max(ys1)),
                }
            )
            current_texts = []
            current_boxes = []
            current_chars = 0

        for paragraph in paragraphs:
            text = str(paragraph.get("text") or "").strip()
            bbox = paragraph.get("bbox")
            if not text or not isinstance(bbox, tuple) or len(bbox) != 4:
                continue

            next_chars = current_chars + (1 if current_texts else 0) + len(text)
            should_flush = current_texts and current_chars >= self.PDF_TARGET_CHUNK_CHARS and next_chars > self.PDF_MAX_CHUNK_CHARS
            if should_flush:
                flush()

            current_texts.append(text)
            current_boxes.append(bbox)
            current_chars += (1 if current_chars else 0) + len(text)

            if current_chars >= self.PDF_MAX_CHUNK_CHARS:
                flush()

        flush()
        return merged


parser = DocumentParser()
